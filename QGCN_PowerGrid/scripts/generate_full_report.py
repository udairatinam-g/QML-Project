import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import chapters_data

def apply_formatting(doc):
    # Setup styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Heading 1 style (for Chapters)
    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Times New Roman'
    font_h1.size = Pt(16)
    font_h1.bold = True
    font_h1.color.rgb = RGBColor(0, 0, 0) # Black
    style_h1.paragraph_format.space_after = Pt(12)
    
    # Heading 2 style
    style_h2 = doc.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Times New Roman'
    font_h2.size = Pt(14)
    font_h2.bold = True
    font_h2.color.rgb = RGBColor(0, 0, 0) # Black
    style_h2.paragraph_format.space_after = Pt(6)

    # Heading 3 style
    try:
        style_h3 = doc.styles['Heading 3']
    except KeyError:
        style_h3 = doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    font_h3 = style_h3.font
    font_h3.name = 'Times New Roman'
    font_h3.size = Pt(12)
    font_h3.bold = True
    font_h3.italic = True
    font_h3.color.rgb = RGBColor(0, 0, 0) # Black
    style_h3.paragraph_format.space_after = Pt(6)
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(0.5)

def add_paragraph(doc, text, style='Normal', align='justify', bold=False, italic=False, size=None):
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.line_spacing = 1.5
    if align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    for run in p.runs:
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if size:
            run.font.size = Pt(size)
    return p

def add_code_block(doc, code_text):
    doc.add_paragraph("\n")
    chunks = code_text.split('\n')
    for chunk in chunks:
        p = doc.add_paragraph(chunk)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
    doc.add_paragraph("\n")

def add_image_with_caption(doc, image_path, caption):
    if os.path.exists(image_path):
        doc.add_paragraph("\n")
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(image_path, width=Inches(5.0))
        p_cap = add_paragraph(doc, caption, align='center', italic=True, size=11)
        doc.add_paragraph("\n")
    else:
        add_paragraph(doc, f"[Image Missing: {image_path}]", align='center', bold=True)

# FRONT MATTER
def add_title_page(doc):
    doc.add_paragraph("\n\n\n")
    add_paragraph(doc, "QUANTUM GRAPH CONVOLUTIONAL NETWORKS FOR POWER GRID RELIABILITY", align='center', bold=True, size=18)
    doc.add_paragraph("\n")
    add_paragraph(doc, "21CSP302L- PROJECT", align='center', bold=True, size=14)
    doc.add_paragraph("\n")
    add_paragraph(doc, "Submitted by", align='center', size=14)
    add_paragraph(doc, "Udai Ratinam G [RA2311003011971]", align='center', italic=True, size=14)
    doc.add_paragraph("\n")
    add_paragraph(doc, "Under the Guidance of", align='center', size=16)
    add_paragraph(doc, "Dr. Nalini S", align='center', italic=True, size=14)
    add_paragraph(doc, "(Assistant Professor, Department of Computing Technologies)", align='center', size=16)
    doc.add_paragraph("\n")
    add_paragraph(doc, "in partial fulfillment of the requirements for the degree of", align='center', size=12)
    doc.add_paragraph("\n")
    add_paragraph(doc, "BACHELOR OF TECHNOLOGY", align='center', bold=True, size=14)
    add_paragraph(doc, "in", align='center', size=14)
    add_paragraph(doc, "COMPUTER SCIENCE ENGINEERING", align='center', bold=True, size=14)
    doc.add_paragraph("\n\n\n")
    add_paragraph(doc, "DEPARTMENT OF COMPUTATIONAL INTELLIGENCE COLLEGE OF ENGINEERING AND TECHNOLOGY", align='center', size=16)
    add_paragraph(doc, "SRM INSTITUTE OF SCIENCE AND TECHNOLOGY", align='center', size=16)
    add_paragraph(doc, "KATTANKULATHUR- 603 203", align='center', size=16)
    add_paragraph(doc, "MAY 2026", align='center', size=14)
    doc.add_page_break()

def add_declaration_page(doc):
    add_paragraph(doc, "Department of Computational Intelligence", align='center', size=14)
    add_paragraph(doc, "SRM Institute of Science & Technology", align='center', size=14)
    add_paragraph(doc, "Own Work Declaration Form", align='center', size=14, bold=True)
    doc.add_paragraph("\n")
    add_paragraph(doc, "This sheet must be filled in (each box ticked to show that the condition has been met). It must be signed and dated along with your student registration number and included with all assignments you submit – work will not be marked unless this is done.", size=12)
    add_paragraph(doc, "To be completed by the student for all assessments", size=12)
    add_paragraph(doc, "Degree/ Course    : B.Tech CSE", size=12)
    add_paragraph(doc, "Student Name      : Udai Ratinam G", size=12)
    add_paragraph(doc, "Registration Number : RA2311003011971", size=12)
    add_paragraph(doc, "Title of Work     : Quantum Graph Convolutional Networks for Power Grid Reliability", size=12)
    doc.add_paragraph("\n")
    add_paragraph(doc, "I hereby certify that this assessment compiles with the University's Rules and Regulations relating to Academic misconduct and plagiarism.", size=12)
    add_paragraph(doc, "I confirm that all the work contained in this assessment is my own except where indicated, and that I have met the following conditions:", size=12)
    add_paragraph(doc, "1. Clearly referenced / listed all sources as appropriate.", size=12)
    add_paragraph(doc, "2. Referenced and put in inverted commas all quoted text.", size=12)
    add_paragraph(doc, "3. Given the sources of all pictures, data etc. that are not my own.", size=12)
    add_paragraph(doc, "4. Not made any use of the report of any other student.", size=12)
    add_paragraph(doc, "5. Acknowledged in appropriate places any help that I have received from others.", size=12)
    doc.add_paragraph("\n\n\n")
    add_paragraph(doc, "Signature: ______________________", align='right', size=12)
    add_paragraph(doc, "Date: ______________________", align='right', size=12)
    doc.add_page_break()

def add_bonafide_certificate(doc):
    add_paragraph(doc, "SRM INSTITUTE OF SCIENCE AND TECHNOLOGY KATTANKULATHUR - 603 203", align='center', size=18, bold=True)
    add_paragraph(doc, "BONAFIDE CERTIFICATE", align='center', size=16, bold=True)
    doc.add_paragraph("\n\n")
    add_paragraph(doc, "Certified that 21CSP302L - Project report titled \"QUANTUM GRAPH CONVOLUTIONAL NETWORKS FOR POWER GRID RELIABILITY\" is the bonafide work of \"Udai Ratinam G [RA2311003011971]\" who carried out the project work under my supervision. Certified further, that to the best of my knowledge the work reported herein does not form any other project report or dissertation on the basis of which a degree or award was conferred on an earlier occasion on this or any other candidate.", size=14)
    doc.add_paragraph("\n\n\n\n")
    add_paragraph(doc, "<<Signature>>                                      <<Signature>>", size=12)
    add_paragraph(doc, "EXAMINER 1                                         EXAMINER 2", size=12)
    add_paragraph(doc, "Name & Signature                                   Name & Signature", size=12)
    doc.add_page_break()

def add_acknowledgements(doc):
    add_paragraph(doc, "ACKNOWLEDGEMENTS", align='center', size=16, bold=True)
    doc.add_paragraph("\n")
    add_paragraph(doc, "We express our humble gratitude to Dr. C. Muthamizhchelvan, Vice-Chancellor, SRM Institute of Science and Technology, for the facilities extended for the project work and his continued support.")
    add_paragraph(doc, "We extend our sincere thanks to Dr. Leenus Jesu Martin M, Dean-CET, SRM Institute of Science and Technology, for his invaluable support.")
    add_paragraph(doc, "Our inexpressible respect and thanks to our guide, Dr. Nalini S, Assistant Professor, Department of Computing Technologies, SRM Institute of Science and Technology, for providing us with an opportunity to pursue our project under her mentorship. She provided us with the freedom and support to explore the research topics of our interest. Her passion for solving problems and making a difference in the world has always been inspiring.")
    add_paragraph(doc, "Finally, we would like to thank our parents, family members, and friends for their unconditional love, constant support and encouragement.")
    doc.add_paragraph("\n\n")
    add_paragraph(doc, "Author", align='right')
    add_paragraph(doc, "Udai Ratinam G", align='right')
    doc.add_page_break()

def add_abstract(doc):
    add_paragraph(doc, "ABSTRACT", align='center', size=16, bold=True)
    doc.add_paragraph("\n")
    text = (
        "Modern power grids represent critical national infrastructure characterized by highly complex, "
        "interconnected graph topologies. Predicting reliability and preventing cascading failures in such "
        "networks is fundamentally a graph learning problem. Traditional classical models often struggle "
        "with the high-dimensional feature spaces of transmission states and the intricate web of nodal "
        "dependencies. This project introduces a novel hybrid approach: Quantum Graph Convolutional Networks (QGCN), "
        "which leverage quantum state embeddings to process smart-grid topological data. "
        "\n\n"
        "The architecture begins by constructing graph representations of grid topologies (e.g., IEEE 14-bus systems) "
        "annotated with physical telemetry such as voltage, phase, and power flows. These features are encoded "
        "into a higher-dimensional Hilbert space using parameterized quantum circuits, allowing the model to "
        "exploit quantum entanglement and superposition to capture complex non-linear correlations efficiently. "
        "A classical message-passing layer aggregates these quantum-enhanced features to predict grid reliability indices "
        "and potential failure states. "
        "\n\n"
        "Through rigorous simulation using Qiskit and PennyLane, we demonstrate that the hybrid QGCN framework successfully learns "
        "to classify hazardous grid conditions. The integration of quantum representation provides theoretical advantages "
        "in feature expressivity. This report details the software architecture, methodology, core code implementation, "
        "and empirical results of applying quantum machine learning to critical infrastructure reliability, offering "
        "a robust pipeline towards quantum-native smart grid management."
    )
    add_paragraph(doc, text, size=14)
    doc.add_page_break()

def add_toc(doc):
    add_paragraph(doc, "TABLE OF CONTENTS", align='center', size=16, bold=True)
    toc_lines = [
        "ABSTRACT\t\t\t\t\t\t\t\tv",
        "TABLE OF CONTENTS\t\t\t\t\t\tvi",
        "LIST OF FIGURES\t\t\t\t\t\t\tvii",
        "\nCHAPTER 1\tINTRODUCTION & SYSTEM ARCHITECTURE\t1",
        "  1.1 Overview of the Software Pipeline\t\t\t\t1",
        "  1.2 Core Dependencies and Simulation Stack\t\t\t2",
        "  1.3 Project Implementation Goals\t\t\t\t2",
        "CHAPTER 2\tDATA MODELING & GRAPH CONSTRUCTION\t4",
        "  2.1 IEEE 14-Bus Topology Instantiation\t\t\t4",
        "  2.2 Telemetry Feature Normalization\t\t\t\t5",
        "CHAPTER 3\tQUANTUM FEATURE ENCODER IMPLEMENTATION\t7",
        "  3.1 PennyLane Simulator Configuration\t\t\t\t7",
        "  3.2 Core Code: Angle Encoding Circuit\t\t\t\t8",
        "  3.3 Core Code: IQP Encoding Architecture\t\t\t10",
        "CHAPTER 4\tCLASSICAL MESSAGE PASSING\t\t\t12",
        "  4.1 Adjacency Matrix Multiplication\t\t\t\t12",
        "  4.2 Hybrid Gradient Computation\t\t\t\t12",
        "CHAPTER 5\tRESULTS & VISUALIZATIONS\t\t\t14",
        "  5.1 Model Convergence Analysis\t\t\t\t14",
        "  5.2 Spatial Risk Visualization\t\t\t\t14",
        "  5.3 Baseline Comparison Analysis\t\t\t\t16",
        "\nREFERENCES\t\t\t\t\t\t\t18"
    ]
    for line in toc_lines:
        add_paragraph(doc, line, size=12)
    doc.add_page_break()

# EXTENSIVE CONTENT GENERATION
def add_chapters(doc, outputs_dir):
    data = chapters_data.implementation_paragraphs

    # CHAPTER 1
    add_paragraph(doc, "CHAPTER 1 - INTRODUCTION & SYSTEM ARCHITECTURE", style='Heading 1', align='center')
    add_paragraph(doc, "1.1 Overview of the Software Pipeline", style='Heading 2')
    add_paragraph(doc, data[0])
    add_paragraph(doc, data[1])
    add_image_with_caption(doc, os.path.join(outputs_dir, "08_architecture_flow.png"), "Figure 1.1: System Architecture and Pipeline Flow")
    
    add_paragraph(doc, "1.2 Core Dependencies and Simulation Stack", style='Heading 2')
    add_paragraph(doc, data[2])
    add_paragraph(doc, data[3])
    add_paragraph(doc, data[4])
    add_paragraph(doc, data[5])
    
    add_paragraph(doc, "1.3 Project Implementation Goals", style='Heading 2')
    add_paragraph(doc, data[6])
    add_paragraph(doc, data[7])
    doc.add_page_break()
    
    # CHAPTER 2
    add_paragraph(doc, "CHAPTER 2 - DATA MODELING & GRAPH CONSTRUCTION", style='Heading 1', align='center')
    add_paragraph(doc, "2.1 IEEE 14-Bus Topology Instantiation", style='Heading 2')
    add_paragraph(doc, data[8])
    add_paragraph(doc, data[9])
    add_paragraph(doc, data[10])
    add_image_with_caption(doc, os.path.join(outputs_dir, "01_grid_topology.png"), "Figure 2.1: IEEE 14-bus Grid Topology Extracted from networkx")
    
    add_paragraph(doc, "2.2 Telemetry Feature Normalization", style='Heading 2')
    add_paragraph(doc, data[11])
    add_paragraph(doc, data[12])
    add_paragraph(doc, data[13])
    add_paragraph(doc, data[14])
    add_paragraph(doc, data[15])
    doc.add_page_break()

    # CHAPTER 3
    add_paragraph(doc, "CHAPTER 3 - QUANTUM FEATURE ENCODER IMPLEMENTATION", style='Heading 1', align='center')
    add_paragraph(doc, "3.1 PennyLane Simulator Configuration", style='Heading 2')
    add_paragraph(doc, data[16])
    add_paragraph(doc, data[17])
    add_paragraph(doc, data[18])
    add_image_with_caption(doc, os.path.join(outputs_dir, "05_quantum_states.png"), "Figure 3.1: Visualization of Quantum States Post-Embedding")
    
    add_paragraph(doc, "3.2 Core Code: Angle Encoding Circuit", style='Heading 2')
    add_paragraph(doc, data[19])
    angle_code = '''
    def angle_encoding(self, features: np.ndarray) -> np.ndarray:
        """
        Angle Encoding: Map feature values to rotation angles.
        Each qubit rotates by an angle proportional to a feature.
        """
        @qml.qnode(self.dev)
        def encoded_circuit(x):
            # Normalize features to [-π, π]
            x_normalized = np.clip(x, -1, 1) * np.pi
            
            # Use non-commuting rotations so Z-basis measurements carry feature signal.
            for i in range(self.n_qubits):
                theta = x_normalized[i % len(x_normalized)]
                phi = x_normalized[(i + 1) % len(x_normalized)] if len(x_normalized) > 1 else theta
                qml.RX(theta, wires=i)
                qml.RY(phi, wires=i)

            if self.n_qubits > 1:
                for i in range(self.n_qubits - 1):
                    qml.CNOT(wires=[i, i + 1])
            
            # Apply X-basis measurement for feature extraction
            return [qml.expval(qml.PauliZ(i)) for i in range(self.n_qubits)]
        
        return np.array(encoded_circuit(features))
    '''
    add_code_block(doc, angle_code)
    add_paragraph(doc, data[20])
    add_paragraph(doc, data[21])
    
    add_paragraph(doc, "3.3 Core Code: IQP Encoding Architecture", style='Heading 2')
    add_paragraph(doc, data[22])
    add_image_with_caption(doc, os.path.join(outputs_dir, "09_quantum_circuit_diagram.png"), "Figure 3.2: IQP Encoding Circuit Abstraction")
    add_paragraph(doc, data[23])
    add_paragraph(doc, data[24])
    add_image_with_caption(doc, os.path.join(outputs_dir, "qkernel_entangled.png"), "Figure 3.3: Entangled Quantum Kernel Matrix Representation")
    doc.add_page_break()
    
    # CHAPTER 4
    add_paragraph(doc, "CHAPTER 4 - CLASSICAL MESSAGE PASSING", style='Heading 1', align='center')
    add_paragraph(doc, "4.1 Adjacency Matrix Multiplication", style='Heading 2')
    add_paragraph(doc, data[25])
    add_paragraph(doc, data[26])
    add_paragraph(doc, data[27])
    add_paragraph(doc, data[28])
    add_paragraph(doc, "4.2 Hybrid Gradient Computation", style='Heading 2')
    add_paragraph(doc, data[29])
    add_paragraph(doc, data[30])
    add_paragraph(doc, data[31])
    add_paragraph(doc, data[32])
    add_paragraph(doc, data[33])
    doc.add_page_break()
    
    # CHAPTER 5
    add_paragraph(doc, "CHAPTER 5 - RESULTS & VISUALIZATIONS", style='Heading 1', align='center')
    add_paragraph(doc, "5.1 Model Convergence Analysis", style='Heading 2')
    add_paragraph(doc, data[34])
    add_image_with_caption(doc, os.path.join(outputs_dir, "02_training_history.png"), "Figure 5.1: Model Training History and Loss Convergence")
    add_paragraph(doc, data[35])
    
    add_paragraph(doc, "5.2 Spatial Risk Visualization", style='Heading 2')
    add_image_with_caption(doc, os.path.join(outputs_dir, "04_grid_risk_heatmap.png"), "Figure 5.2: High-fidelity Grid Risk Heatmap")
    add_paragraph(doc, data[36])
    add_image_with_caption(doc, os.path.join(outputs_dir, "03_prediction_distribution.png"), "Figure 5.3: Prediction Distribution Density")
    add_paragraph(doc, data[37])
    
    add_paragraph(doc, "5.3 Baseline Comparison Analysis", style='Heading 2')
    add_paragraph(doc, data[38])
    add_image_with_caption(doc, os.path.join(outputs_dir, "06_roc_curve.png"), "Figure 5.4: Receiver Operating Characteristic (ROC)")
    add_paragraph(doc, data[39])
    add_image_with_caption(doc, os.path.join(outputs_dir, "07_baseline_comparison.png"), "Figure 5.5: Classical GCN vs Quantum GCN Baselines")
    add_paragraph(doc, data[40])
    add_paragraph(doc, data[41])
    add_paragraph(doc, data[42])
    doc.add_page_break()

def add_references(doc):
    add_paragraph(doc, "REFERENCES", align='center', size=16, bold=True)
    refs = [
        "[1] M. Schuld and F. Petruccione, Supervised Learning with Quantum Computers. Springer, 2018.",
        "[2] T. N. Kipf and M. Welling, \"Semi-Supervised Classification with Graph Convolutional Networks,\" ICLR, 2017.",
        "[3] Qiskit contributors, \"Qiskit: An Open-source Framework for Quantum Computing,\" 2023.",
        "[4] V. Havlíček et al., \"Supervised learning with quantum-enhanced feature spaces,\" Nature, vol. 567, pp. 209-212, 2019.",
        "[5] IEEE Power & Energy Society, \"IEEE 14-bus system topology and telemetry standards,\" IEEE, 2020.",
        "[6] A. Abbas et al., \"The power of quantum neural networks,\" Nature Computational Science, vol. 1, pp. 403-409, 2021."
    ]
    for ref in refs:
        add_paragraph(doc, ref)
    doc.add_page_break()

def main():
    doc = Document()
    apply_formatting(doc)
    
    add_title_page(doc)
    add_declaration_page(doc)
    add_bonafide_certificate(doc)
    add_acknowledgements(doc)
    add_abstract(doc)
    add_toc(doc)
    
    outputs_dir = r"C:\Users\Udai Ratinam G\Downloads\QML Project\QGCN_PowerGrid\outputs"
    add_chapters(doc, outputs_dir)
    add_references(doc)
    
    out_dir = Path(r"C:\Users\Udai Ratinam G\Downloads\QML Project\QGCN_PowerGrid\deliverables")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "QGCN_Formal_Project_Report_v7.docx"
    doc.save(str(out_path))
    print(f"Successfully generated {out_path}")

if __name__ == "__main__":
    main()
