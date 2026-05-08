from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime


def build_abstract() -> None:
    doc = Document()

    # Keep margins and typography consistent with the project write-up style.
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ABSTRACT")
    title_run.bold = True
    title_run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Quantum Graph Convolutional Networks for Power Grid Reliability")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Udai Ratinam G | Batch 2 | Reg No: RA2311003011971")

    doc.add_paragraph()

    abstract_text = (
        "Modern power grids are complex graph-structured systems in which local disturbances can "
        "propagate across connected substations and transmission lines. This work presents a hybrid "
        "Quantum Graph Convolutional Network (QGCN) framework for power-grid reliability modeling, "
        "with the aim of learning risk patterns directly from node features and network topology. "
        "Unlike conventional learning pipelines that process each operating point independently, the "
        "proposed method preserves structural dependencies and captures interaction effects among "
        "electrically connected buses.\n\n"
        "The framework combines graph learning with variational quantum processing. Grid observations "
        "are first represented as node-level feature vectors, including voltage and power-state "
        "attributes relevant to operational stability. These features are then embedded into quantum "
        "states through parameterized encoding, followed by a trainable quantum layer with "
        "entanglement to model non-linear correlations that are difficult to capture using purely "
        "classical architectures. The learned quantum representations are integrated with a classical "
        "decision stage to infer reliability-oriented risk scores for nodes and the overall grid "
        "state.\n\n"
        "This study contributes an application-focused formulation of quantum-enhanced graph learning "
        "for critical infrastructure analytics, demonstrating how QGCN concepts can be adapted to "
        "power-system reliability assessment in a practical workflow. The proposed approach supports "
        "interpretable early-warning analysis and provides a foundation for future extensions toward "
        "larger topologies, uncertainty-aware forecasting, and integration with real supervisory "
        "monitoring data.\n\n"
        "Beyond immediate reliability prediction, the study highlights a broader methodological path "
        "for combining emerging quantum techniques with domain-constrained engineering models. By "
        "retaining physically meaningful grid structure while enriching representation capacity through "
        "quantum feature transformation, the framework encourages a balanced design philosophy in "
        "which interpretability, operational relevance, and computational innovation are treated as "
        "co-equal objectives. This perspective is especially valuable for safety-critical infrastructure, "
        "where decision support systems must not only be accurate but also transparent, adaptable, and "
        "aligned with real-time monitoring and control practices."
    )

    doc.add_paragraph(abstract_text)

    deliverables_dir = Path("deliverables")
    deliverables_dir.mkdir(parents=True, exist_ok=True)
    output_path = deliverables_dir / "QGCN_Abstract_One_Page_Udai_Ratinam_G.docx"
    try:
        doc.save(str(output_path))
        print(f"Created: {output_path}")
    except PermissionError:
        fallback_path = (
            deliverables_dir / "QGCN_Abstract_One_Page_Udai_Ratinam_G_Updated.docx"
        )
        try:
            doc.save(str(fallback_path))
            print(f"Created: {fallback_path}")
        except PermissionError:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            timestamped_path = (
                deliverables_dir
                / f"QGCN_Abstract_One_Page_Udai_Ratinam_G_{timestamp}.docx"
            )
            doc.save(str(timestamped_path))
            print(f"Created: {timestamped_path}")


if __name__ == "__main__":
    build_abstract()
