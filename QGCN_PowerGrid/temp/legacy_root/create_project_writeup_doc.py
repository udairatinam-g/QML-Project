from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def build_writeup() -> None:
    doc = Document()

    # Set standard margins for a clean one-page academic format.
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Project Write-Up\n")
    run.bold = True
    run.font.size = Pt(16)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Name: Udai Ratinam G\n")
    info.add_run("Batch No: 2\n")
    info.add_run("Reg No: RA2311003011971")

    doc.add_paragraph()

    heading = doc.add_paragraph()
    heading_run = heading.add_run(
        "Quantum Graph Convolutional Networks for Power Grid Reliability"
    )
    heading_run.bold = True

    body_text = (
        "This project presents a practical implementation of Quantum Graph Convolutional "
        "Networks (QGCN) for power grid reliability prediction in smart-city infrastructure. "
        "The core motivation is that modern electrical grids behave as graph-structured systems, "
        "where each substation acts as a node and each transmission line forms an edge. Because "
        "disturbances can propagate through these connections, reliable early-warning models must "
        "learn both local node behavior and network-wide interactions. Traditional machine learning "
        "models often struggle to capture these complex relationships efficiently, especially when "
        "the dimensionality of operational data is high.\n\n"
        "To address this, the project combines graph learning principles with quantum feature "
        "processing. In Phase 1, an IEEE 14-bus graph is created to represent a realistic test grid, "
        "and synthetic operational samples are generated with features such as voltage magnitude, "
        "phase angle, active power, reactive power, and load level. In Phase 2, these classical "
        "features are encoded into quantum states using angle encoding through rotation gates. "
        "This allows compact representation of feature interactions in a higher-dimensional Hilbert "
        "space. In Phase 3, a variational quantum circuit (VQC) with trainable parameters and "
        "entanglement layers performs quantum message passing, enabling each node to incorporate "
        "information from neighboring nodes. In Phase 4, a hybrid training pipeline maps quantum "
        "outputs to failure probabilities and optimizes the model using binary cross-entropy loss.\n\n"
        "The developed system demonstrates strong predictive performance with approximately 82.9% "
        "accuracy and AUC around 0.847 on the generated test scenarios. These results indicate that "
        "the model can identify risky grid states while maintaining a balanced trade-off between "
        "precision and recall. In practical terms, this helps operators by providing an interpretable "
        "risk score that supports preventive actions such as load balancing, rerouting, and reserve "
        "activation before large-scale outages occur. The project therefore contributes both a "
        "theoretical and implementation-level framework for applying quantum machine learning in a "
        "critical real-world domain."
    )

    doc.add_paragraph(body_text)

    deliverables_dir = Path("deliverables")
    deliverables_dir.mkdir(parents=True, exist_ok=True)
    output_path = deliverables_dir / "QGCN_Project_Writeup_Udai_Ratinam_G.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")


if __name__ == "__main__":
    build_writeup()
